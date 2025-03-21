# simple send the mails 
'''
import smtplib
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
import docx

# Function to read emails from the docx file
def read_emails_from_docx(file_path):
    doc = docx.Document(file_path)
    emails = []
    email = {}
    for para in doc.paragraphs:
        if para.text.startswith("Sample Email"):
            if email:
                emails.append(email)
                email = {}
        if para.text.startswith("Subject:"):
            email['subject'] = para.text.replace("Subject: ", "")
        elif para.text.startswith("Email:"):
            email['body'] = ""
        else:
            if 'body' in email:
                email['body'] += para.text + "\n"
    if email:
        emails.append(email)
    return emails

# Function to send an email
def send_email(subject, body, recipient_email, sender_email, sender_password):
    msg = MIMEMultipart()
    msg['From'] = sender_email
    msg['To'] = recipient_email
    msg['Subject'] = subject

    msg.attach(MIMEText(body, 'plain'))

    try:
        server = smtplib.SMTP('smtp.gmail.com', 587)
        server.starttls()
        server.login(sender_email, sender_password)
        text = msg.as_string()
        server.sendmail(sender_email, recipient_email, text)
        server.quit()
        print(f"Email sent successfully: {subject}")
    except Exception as e:
        print(f"Failed to send email: {subject}. Error: {str(e)}")

# Main function to read emails from the docx file and send them one by one
def main():
    file_path = 'e:/one drive/OneDrive - Happy English/programms/new customer care/customer-care-system/core data/Sample Email to test.docx'
    #recipient_email = 'fa22-bse-031@cuivehari.edu.pk'  # Replace with the recipient's email address
    recipient_email = 'saqlainfawad@gmail.com'
    sender_email = 'saqlainfawad@gmail.com'  # Replace with your email address
    sender_password = 'jkgc keje xjmz xqyp'  # Replace with your email password

    emails = read_emails_from_docx(file_path)
    
    for email in emails:
        send_email(email['subject'], email['body'], recipient_email, sender_email, sender_password)

if __name__ == "__main__":
    main()
'''

import smtplib
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
import docx
import random

# Function to read emails from the docx file and add random data
def read_emails_from_docx(file_path):
    doc = docx.Document(file_path)
    emails = []
    email = {}
    for para in doc.paragraphs:
        if para.text.startswith("Sample Email"):
            if email:
                emails.append(email)
                email = {}
        if para.text.startswith("Subject:"):
            email['subject'] = para.text.replace("Subject: ", "")
        elif para.text.startswith("Email:"):
            email['body'] = ""
        else:
            if 'body' in email:
                email['body'] += para.text + "\n"
    if email:
        emails.append(email)
    return emails

# Function to add random data to the placeholders
def add_random_data(emails):
    recipient_names = ["John Doe", "Jane Smith", "Michael Johnson", "Emily Davis"]
    sender_names = ["Alice Brown", "David Wilson", "Sophia Martinez", "James Anderson"]
    positions = ["Customer Service Representative", "Sales Manager", "Support Specialist", "Account Executive"]
    contact_info = ["alice.brown@example.com", "david.wilson@example.com", "sophia.martinez@example.com", "james.anderson@example.com"]
    company_names = ["Tech Solutions Inc.", "Global Enterprises", "Innovative Tech Co.", "Business Corp"]

    for email in emails:
        email['body'] = email['body'].replace("[Recipient's Name]", random.choice(recipient_names))
        email['body'] = email['body'].replace("[Your Name]", random.choice(sender_names))
        email['body'] = email['body'].replace("[Your Position]", random.choice(positions))
        email['body'] = email['body'].replace("[Your Contact Information]", random.choice(contact_info))
        email['body'] = email['body'].replace("[Your Company Name]", random.choice(company_names))

    return emails

# Function to send an email
def send_email(subject, body, recipient_email, sender_email, sender_password):
    msg = MIMEMultipart()
    msg['From'] = sender_email
    msg['To'] = recipient_email
    msg['Subject'] = subject

    msg.attach(MIMEText(body, 'plain'))

    try:
        server = smtplib.SMTP('smtp.gmail.com', 587)
        server.starttls()
        server.login(sender_email, sender_password)
        text = msg.as_string()
        server.sendmail(sender_email, recipient_email, text)
        server.quit()
        print(f"Email sent successfully: {subject}")
    except Exception as e:
        print(f"Failed to send email: {subject}. Error: {str(e)}")

# Main function to read emails from the docx file, add random data, and send them one by one
def main():
    file_path = 'e:/one drive/OneDrive - Happy English/programms/new customer care/customer-care-system/core data/Sample Email to test.docx'
    recipient_email = 'saqlainfawad@gmail.com'  # Replace with the recipient's email address
    sender_email = 'saqlainfawad@gmail.com'  # Replace with your email address
    sender_password = 'jkgc keje xjmz xqyp'  # Replace with your email password

    emails = read_emails_from_docx(file_path)
    updated_emails = add_random_data(emails)
    
    for email in updated_emails:
        send_email(email['subject'], email['body'], recipient_email, sender_email, sender_password)

if __name__ == "__main__":
    main()